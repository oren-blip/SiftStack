#!/usr/bin/env python3
"""
personalize.py — Deterministic personalization for the DataSift Lead Management Training.

FILE_VERSION = "0.1.8" — if you're reading this and the EOF marker at the bottom
of this file is missing, your copy is truncated. Re-download the .plugin file.

Reads .datasift-config.json from the working folder and produces TWO personalized files:
  1. Lead-Manager-Training.html  (find/replace on the bundled brand-standard HTML template)
  2. Lead-Manager-Training-Manual.docx  (find/replace on the bundled docx template)

Uses two-pass sentinel substitution so company names that contain substrings of the
original template never cascade. Also handles lead source personalization (website,
PPL, direct mail inboxes).

Usage:
    python3 personalize.py <working_folder>
"""
import json
import sys
import os
import shutil
from pathlib import Path

# ── Two-pass sentinel substitution ──────────────────────────────────────────

def build_rules(cfg):
    """Build (patterns, sentinels) for two-pass substitution."""
    company = cfg.get("company_name", "").strip() or "Your Company"
    user = cfg.get("user_name", "").strip() or "Your Team"
    phone = cfg.get("phone_provider", "").strip() or "your phone provider"
    company_lower = company.lower()

    # Lead source config (with defaults matching the FCRE template)
    ls = cfg.get("lead_sources", {})

    ws = ls.get("website", {})
    ws_active = ws.get("active", False)
    ws_platform = (ws.get("platform") or "").strip() or "Your Website"
    ws_inbox = (ws.get("inbox") or "").strip() or "Online Inbounds"

    ppl = ls.get("ppl", {})
    ppl_active = ppl.get("active", False)
    ppl_provider = (ppl.get("provider") or "").strip() or "your PPL provider"
    ppl_inbox = (ppl.get("inbox") or "").strip() or "PPL Inbounds"

    dm = ls.get("direct_mail", {})
    dm_inbox = (dm.get("inbox") or "").strip() or "Inbound Leads"

    # Build the pillar header replacements
    # If a source is inactive, we'll add [Optional] to the header
    website_header = f"{ws_platform} (Website)" if ws_active else f"{ws_platform} (Website) <span class=\"ds-optional-badge\">[Optional]</span>"
    ppl_header = f"Pay-Per-Lead ({ppl_provider})" if ppl_active else f"Pay-Per-Lead <span class=\"ds-optional-badge\">[Optional]</span>"

    # Sentinels — null bytes never appear in HTML/docx source
    sentinels = {
        # Company / user
        "\x00S_USER_WITH_CO\x00": f"{user} with {company}",
        "\x00S_USER_FROM_CO\x00": f"{user} from {company}",
        "\x00S_THIS_USER_FROM\x00": f"This is {user} from {company}",
        "\x00S_CO_KPIS_SHEET\x00": f"{company} KPIs Google Sheet",
        "\x00S_CO_KPIS\x00": f"{company} KPIs",
        "\x00S_CO_TEAM\x00": f"{company} team",
        "\x00S_CO_POSS\x00": f"{company}'s",
        "\x00S_CO\x00": company,
        "\x00S_CO_LO\x00": company_lower,
        # Lead sources
        "\x00S_WS_HEADER\x00": website_header,
        "\x00S_WS_PLATFORM\x00": ws_platform,
        "\x00S_WS_INBOX\x00": ws_inbox,
        "\x00S_PPL_HEADER\x00": ppl_header,
        "\x00S_PPL_PROVIDER\x00": ppl_provider,
        "\x00S_PPL_INBOX\x00": ppl_inbox,
        "\x00S_DM_INBOX\x00": dm_inbox,
        # Phone provider
        "\x00S_PHONE\x00": phone,
    }

    # Patterns — ORDERED LONGEST → SHORTEST
    patterns = [
        # SMS template specifics (most specific)
        ("Tyler with Florida Cash Real Estate", "\x00S_USER_WITH_CO\x00"),
        ("This is Tyler from Florida Cash Real Estate", "\x00S_THIS_USER_FROM\x00"),
        ("Tyler from Florida Cash Real Estate", "\x00S_USER_FROM_CO\x00"),
        ("This is Tyler from FCRE", "\x00S_THIS_USER_FROM\x00"),
        ("Tyler from Florida Cash", "\x00S_USER_FROM_CO\x00"),
        ("Tyler from FCRE", "\x00S_USER_FROM_CO\x00"),
        # Company (longest first)
        ("FCRE KPIs Google Sheet", "\x00S_CO_KPIS_SHEET\x00"),
        ("FCRE KPIs", "\x00S_CO_KPIS\x00"),
        ("FCRE team", "\x00S_CO_TEAM\x00"),
        ("FCRE's", "\x00S_CO_POSS\x00"),
        ("Florida Cash Real Estate", "\x00S_CO\x00"),
        ("florida cash real estate", "\x00S_CO_LO\x00"),
        ("Florida Cash", "\x00S_CO\x00"),
        ("FCRE", "\x00S_CO\x00"),
        # Lead sources — website/Carrot (longest first)
        ('Carrot leads (carrot.com + smrtPhone "Online Inbounds")',
         f'\x00S_WS_PLATFORM\x00 leads (\x00S_WS_PLATFORM\x00 + \x00S_PHONE\x00 "\x00S_WS_INBOX\x00")'),
        ('Process through carrot.com and the smrtPhone "Online Inbounds" inbox',
         f'Process through \x00S_WS_PLATFORM\x00 and the \x00S_PHONE\x00 "\x00S_WS_INBOX\x00" inbox'),
        ("Carrot (Website)", "\x00S_WS_HEADER\x00"),
        ("inbound Carrot lead", f"inbound \x00S_WS_PLATFORM\x00 lead"),
        ("PPC/SEO/Carrot lead", f"PPC/SEO/\x00S_WS_PLATFORM\x00 lead"),
        ("Carrot lead", f"\x00S_WS_PLATFORM\x00 lead"),
        ("Carrot Lead Processing (Part 2)", f"\x00S_WS_PLATFORM\x00 Lead Processing (Part 2)"),
        ("Carrot Lead Processing", f"\x00S_WS_PLATFORM\x00 Lead Processing"),
        ("Carrot Lead with Property Typo", f"\x00S_WS_PLATFORM\x00 Lead with Property Typo"),
        # Lead sources — PPL (longest first)
        ('PPL leads (propertyleads.com + smrtPhone "PPL Inbounds")',
         f'\x00S_PPL_PROVIDER\x00 leads (\x00S_PPL_PROVIDER\x00 + \x00S_PHONE\x00 "\x00S_PPL_INBOX\x00")'),
        ('Process through propertyleads.com and the smrtPhone "PPL Inbounds" inbox',
         f'Process through \x00S_PPL_PROVIDER\x00 and the \x00S_PHONE\x00 "\x00S_PPL_INBOX\x00" inbox'),
        ("Pay-Per-Lead (PPL)", "\x00S_PPL_HEADER\x00"),
        ("propertyleads.com", "\x00S_PPL_PROVIDER\x00"),
        # Lead sources — direct mail inbox
        ('smrtPhone "Inbound Leads"', f'\x00S_PHONE\x00 "\x00S_DM_INBOX\x00"'),
        ("smrtPhone inbound inbox", f"\x00S_PHONE\x00 \"\x00S_DM_INBOX\x00\""),
        # Phone provider (catch-all for remaining smrtPhone refs in inbox contexts)
        ('smrtPhone "Online Inbounds"', f'\x00S_PHONE\x00 "\x00S_WS_INBOX\x00"'),
        ('smrtPhone "PPL Inbounds"', f'\x00S_PHONE\x00 "\x00S_PPL_INBOX\x00"'),
    ]
    return patterns, sentinels


def personalize_text(text, cfg):
    """Two-pass find/replace. Case-sensitive."""
    patterns, sentinels = build_rules(cfg)
    # Pass 1: original strings → sentinels
    for find, sentinel in patterns:
        text = text.replace(find, sentinel)
    # Pass 2: sentinels → final values
    for sentinel, value in sentinels.items():
        text = text.replace(sentinel, value)
    return text


def count_replacements(text, cfg):
    patterns, _ = build_rules(cfg)
    return sum(text.count(f) for f, _ in patterns if f in text)


# ── Optional badge CSS injection ────────────────────────────────────────────

OPTIONAL_BADGE_CSS = """
    .ds-optional-badge {
      display: inline-block;
      font-size: var(--text-xs);
      font-weight: 600;
      color: var(--neutral-500);
      background: var(--neutral-100);
      padding: 2px 8px;
      border-radius: var(--radius-sm);
      margin-left: var(--space-2);
      vertical-align: middle;
      opacity: 0.8;
    }"""


def inject_optional_badge_css(html_text, cfg):
    """If any lead source is marked inactive, inject the [Optional] badge CSS."""
    ls = cfg.get("lead_sources", {})
    ws_active = ls.get("website", {}).get("active", True)
    ppl_active = ls.get("ppl", {}).get("active", True)

    if not ws_active or not ppl_active:
        # Inject CSS before the closing </style> tag
        if "</style>" in html_text:
            html_text = html_text.replace("</style>", OPTIONAL_BADGE_CSS + "\n  </style>", 1)
    return html_text


# ── HTML personalization ────────────────────────────────────────────────────

def personalize_html(template_path, output_path, cfg):
    original = template_path.read_text(encoding="utf-8")
    n = count_replacements(original, cfg)
    personalized = personalize_text(original, cfg)
    personalized = inject_optional_badge_css(personalized, cfg)
    output_path.write_text(personalized, encoding="utf-8")
    return n


# ── DOCX personalization ───────────────────────────────────────────────────

def personalize_docx(template_path, output_path, cfg):
    try:
        from docx import Document
    except ImportError:
        print("ERROR: python-docx not installed. Run: pip3 install python-docx --break-system-packages", file=sys.stderr)
        sys.exit(1)

    shutil.copy(template_path, output_path)
    doc = Document(str(output_path))

    def replace_in_runs(runs):
        if not runs:
            return 0
        full_text = "".join(r.text for r in runs)
        new_text = personalize_text(full_text, cfg)
        if new_text == full_text:
            return 0
        runs[0].text = new_text
        for r in runs[1:]:
            r.text = ""
        return 1

    changes = 0
    for para in doc.paragraphs:
        changes += replace_in_runs(para.runs)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    changes += replace_in_runs(para.runs)
    for section in doc.sections:
        for header in (section.header, section.first_page_header, section.even_page_header):
            if header is None: continue
            for para in header.paragraphs:
                changes += replace_in_runs(para.runs)
        for footer in (section.footer, section.first_page_footer, section.even_page_footer):
            if footer is None: continue
            for para in footer.paragraphs:
                changes += replace_in_runs(para.runs)

    doc.save(str(output_path))
    return changes


# ── Main ───────────────────────────────────────────────────────────────────

def main():
    if len(sys.argv) != 2:
        print("Usage: python3 personalize.py <working_folder>", file=sys.stderr)
        sys.exit(2)

    working = Path(sys.argv[1]).resolve()
    if not working.exists():
        print(f"ERROR: working folder not found: {working}", file=sys.stderr)
        sys.exit(2)

    config_path = working / ".datasift-config.json"
    if not config_path.exists():
        print(f"ERROR: .datasift-config.json not found in {working}", file=sys.stderr)
        print("       Run the setup-my-training skill first.", file=sys.stderr)
        sys.exit(2)

    cfg = json.loads(config_path.read_text(encoding="utf-8"))
    if not cfg.get("company_name"):
        print("ERROR: company_name missing from .datasift-config.json", file=sys.stderr)
        sys.exit(2)

    print(f"Personalizing for: {cfg['company_name']}")

    # Show lead source summary
    ls = cfg.get("lead_sources", {})
    for src_name, src_key in [("Website", "website"), ("PPL", "ppl"), ("Direct Mail", "direct_mail"), ("Cold Calling", "cold_calling")]:
        src = ls.get(src_key, {})
        active = src.get("active", False)
        print(f"  {src_name}: {'active' if active else 'optional (will badge as [Optional])'}")

    script_dir = Path(__file__).parent.resolve()

    # 1) HTML
    html_template = script_dir / "Lead-Manager-Training-Template.html"
    if not html_template.exists():
        print(f"ERROR: HTML template not found at {html_template}", file=sys.stderr)
        sys.exit(2)
    html_out = working / "Lead-Manager-Training.html"
    print("→ Personalizing HTML training...")
    n = personalize_html(html_template, html_out, cfg)
    print(f"    Wrote: {html_out} ({n} text replacements)")

    # 2) DOCX
    docx_template = script_dir / "Lead-Manager-Training-Manual-Template.docx"
    if not docx_template.exists():
        print(f"WARNING: docx template not found at {docx_template}", file=sys.stderr)
    else:
        docx_out = working / "Lead-Manager-Training-Manual.docx"
        print("→ Personalizing DOCX manual...")
        n = personalize_docx(docx_template, docx_out, cfg)
        print(f"    Wrote: {docx_out} ({n} paragraphs modified)")

    print("Done.")


if __name__ == "__main__":
    main()

# ── EOF INTEGRITY MARKER ────────────────────────────────────
# If this line is missing, your copy of personalize.py is truncated.
# Re-download the datasift-lead-management.plugin file.
_FILE_COMPLETE = True
